#!/usr/bin/env python3
"""Generate the canonical US Signal Celar AI Architecture Package v2.0.

Outputs a Word document, two PDFs, two SVG diagrams, two PNG diagrams, and
SHA256SUMS.txt. The generator uses only the supplied US Signal logo.
"""
from __future__ import annotations

import base64
import hashlib
import html
import re
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A3, LETTER, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import BaseDocTemplate, Frame, Image as RLImage, PageBreak, PageTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfgen import canvas
import cairosvg

VERSION = "2.0"
DOC_TITLE = "Celar AI Private Intelligence Architecture"
CLASSIFICATION = "US Signal Internal — Confidential"
OWNER = "US Signal"
PRODUCT = "Pulse"
NAVY = "#072D59"
BLUE = "#009FE3"
MID_BLUE = "#1769AA"
GREEN = "#67B446"
LIGHT_BLUE = "#EAF5FB"
LIGHT_GRAY = "#F2F5F7"
MED_GRAY = "#D4DCE3"
DARK_GRAY = "#34495E"
WHITE = "#FFFFFF"
BLACK = "#17212B"

ROOT = Path(__file__).resolve().parent
OUT = ROOT.parent
LOGO = ROOT / "US_Signal_Logo.jpg"

FILES = {
    "docx": "US_Signal_Celar_AI_Private_Intelligence_Architecture_v2.0.docx",
    "pdf": "US_Signal_Celar_AI_Private_Intelligence_Architecture_v2.0.pdf",
    "diagrams_pdf": "US_Signal_Celar_AI_Architecture_Diagrams_v2.0.pdf",
    "logical_png": "US_Signal_Celar_AI_Logical_Architecture_v2.0.png",
    "logical_svg": "US_Signal_Celar_AI_Logical_Architecture_v2.0.svg",
    "deployment_png": "US_Signal_Celar_AI_Deployment_Network_Architecture_v2.0.png",
    "deployment_svg": "US_Signal_Celar_AI_Deployment_Network_Architecture_v2.0.svg",
}

SECTIONS = [
    ("1. Executive Summary", [
    "Celar AI is the unified operational intelligence system for the US Signal Solution Provider division. It was conceived and engineered under the direction of Dr. Ahmed Adeyemi, Manager of Professional Services, as the central intersection where consulting teams convene, collaborate, and exchange critical project information.",
    "The name draws from celeritas—the Latin concept of swiftness or speed—and from the conventional symbol c for the speed of light in E=mc². This connection honors US Signal's fiber-network heritage while defining the Professional Services mission to translate the speed of light into the speed of delivery.",
    "Celar AI emerged from the operational friction experienced with legacy Changepoint workflows: siloed information, rigid navigation, repetitive administration, slow SOW preparation, fragmented task handoffs, time-entry burden, and delayed financial visibility.",
    "From a solution-provider perspective, Celar AI unifies authorized documents, live system data, delivery workflows, troubleshooting evidence, reports, and financial context so teams can scope, execute, troubleshoot, report, and invoice work faster without abandoning security, governance, source-system ownership, or human accountability.",
        "Celar AI is the private, permission-aware intelligence layer for Pulse. It combines authorized internal documents, governed live system data, deterministic calculations, and a private language model so users receive detailed, source-grounded answers and reviewable drafts without exposing confidential information to public external models.",
        "Every request begins with authentication and authorization. Pulse retrieves only the document sections and live records the effective user is allowed to access. The private Celar AI model reasons over that context, produces a structured result, and measures confidence, source coverage, freshness, conflicts, and calculation validity.",
        "When private evidence is sufficient, the result proceeds directly to private verification. When generic reasoning could improve the response and policy allows it, Pulse creates a minimal sanitized problem capsule, routes it through Module 064 to an approved external provider, and re-verifies the external result against private authoritative evidence before returning it to the user.",
    ]),
    ("2. Scope, Objectives, and Outcomes", [
        "In scope are Pulse web and API services; identity, RBAC and View-As; project, customer, team, record, and field scope; private document retrieval; governed read-only tools; private embeddings and vector search; private model inference; confidence assessment; evidence verification; Module 064 controlled egress; DLP, audit, evaluation, model registry, and training governance.",
        "Out of scope are autonomous approval or submission, unrestricted model-generated SQL, direct model access to production credentials, raw internal-document transmission to public external models, and autonomous changes to prompts, permissions, model weights, deployments, or production routes.",
        "The target outcome is one reusable intelligence fabric for Timesheet, Help and Search, FlowHive, reports, financial analysis, and future authorized Pulse use cases.",
    ]),
    ("3. Architecture Principles and Decisions", [
        "Private first: restricted data is retrieved and reasoned over inside the approved Pulse trust boundary.",
        "Authorization before retrieval: the application—not the model—decides which data may be accessed.",
        "Facts from tools, explanation from models: deterministic services own calculations, permissions, schedules, and record state.",
        "Progressive disclosure: responses lead with a direct conclusion and then expose evidence, calculations, assumptions, conflicts, limitations, risks, and next actions.",
        "External models are optional: Claude or OpenAI may provide generic reasoning only after DLP sanitization and policy approval.",
        "Human control: timesheets, project baselines, financial actions, datasets, model versions, and production promotion remain human reviewed.",
    ]),
    ("4. Logical Architecture", [
        "The logical design separates the Pulse application trust zone, private data and AI zones, controlled egress, and external-provider boundary. Solid lines represent primary operational flow. Dashed lines represent optional sanitized reasoning, telemetry, or audit flow.",
        "The confidence service is the decision point between a complete private answer and optional generic external assistance. External assistance never bypasses private evidence verification.",
    ]),
    ("5. Component Architecture", [
        "Pulse experience layer: Timesheet, Ask Celar AI, global Search, FlowHive, Reporting, Financial Insights, and future module surfaces.",
        "Identity and policy layer: session validation, actual/effective user resolution, module/action permissions, project and customer scope, field restrictions, and data classification.",
        "Knowledge layer: malware-scanned private storage, document versioning, extraction, OCR when required, classification, chunking, embeddings, and permission-filtered hybrid retrieval.",
        "Tool layer: allowlisted read-only module APIs and semantic metric contracts for projects, time, approvals, utilization, capacity, reports, contracts, rates, expenses, billing, invoices, operations, and audit.",
        "Reasoning layer: private model endpoint, prompt and schema registry, confidence service, citation builder, response composer, private verifier, and governed local fallback.",
        "External control layer: sanitization and DLP gateway, Module 064 provider configuration and routing, external provider adapters, circuit breakers, and sanitized telemetry.",
    ]),
    ("6. Trust Zones and Data Boundaries", [
        "Zone A — User and edge: browser, identity provider, WAF/reverse proxy, and authenticated Pulse session.",
        "Zone B — Pulse application: web/API services, authorization, orchestration, module tools, and audit correlation.",
        "Zone C — Private data: PostgreSQL, object storage, extracted text, vector index, retention, and document metadata.",
        "Zone D — Private AI: embedding endpoint, private open-weight model, evaluation runner, model registry, and training jobs.",
        "Zone E — Controlled egress: DLP/redaction, policy approval, Module 064, outbound allowlists, and provider telemetry.",
        "Zone F — External provider: Claude, OpenAI, or another approved provider receiving only a sanitized reasoning capsule.",
    ]),
    ("7. End-to-End Request Lifecycle", [
        "1. Accept a natural-language request or feature-specific generation request from Pulse.",
        "2. Validate the session and resolve actual and effective identities.",
        "3. Resolve module, action, project, customer, team, record, and field scope.",
        "4. Classify the request and build a minimum-necessary retrieval and tool plan.",
        "5. Retrieve authorized document chunks and execute governed read-only tools.",
        "6. Generate a structured private draft using approved prompts and response schemas.",
        "7. Assess confidence, source coverage, freshness, conflicts, and calculation validity.",
        "8. Optionally create a sanitized capsule and route it through Module 064.",
        "9. Re-ground and verify all externally assisted reasoning inside the private boundary.",
        "10. Return a detailed cited answer or reviewable draft and record sanitized audit evidence.",
    ]),
    ("8. Use Case Architecture", [
        "Timesheet: the Engineer’s rough note remains the primary statement of work. Pulse retrieves authorized SOW, GSD, task, request, and engineering documents, drafts a description, and requires the Engineer to review and apply it. Celar AI cannot change hours, save, submit, or approve time.",
        "Help and Search: product Help uses Module 999, approved documentation, the module catalog, workflows, and permission explanations. Live Search uses permission-filtered tools for current records and returns filters, counts, freshness, unknowns, and navigation targets.",
        "FlowHive: Pulse privately extracts scope, deliverables, exclusions, responsibilities, prerequisites, acceptance criteria, quantities, risks, assumptions, dependencies, milestones, and open questions. The deterministic FlowHive engine calculates working dates, critical path, and float. PM and Engineering review remain mandatory.",
        "Reports and Financials: governed semantic tools calculate values and preserve unknowns. The private model explains results, drivers, exceptions, and recommended follow-up. Pulse cannot change rates, expenses, invoices, contracts, or reconciliation state.",
    ]),
    ("9. Private Document Intelligence and RAG", [
        "Documents are malware scanned, classified, versioned, checksummed, approved, and stored privately before extraction. OCR is used only when native text extraction is not sufficient.",
        "Chunks carry security and citation metadata including document ID, project, customer, category, version, classification, visibility, owner, effective date, page or section, checksum, and authorized role/team/user scope.",
        "Hybrid retrieval combines exact keyword matching for codes, model numbers, dates, and contract terms with semantic retrieval for natural-language concepts.",
        "Authorization filters apply before ranking and before prompt assembly. Current approved versions are preferred while conflicting versions are surfaced explicitly.",
    ]),
    ("10. Governed Live-Data Tools and Semantic Layer", [
        "Celar AI never receives unrestricted database credentials. Information is exposed through allowlisted read-only APIs and semantic metric contracts owned by the applicable Pulse module.",
        "The tool gateway validates the requested metrics, dimensions, filters, maximum rows, purpose, sensitivity, module access, and record scope before execution.",
        "Financial and operational answers show formula or metric definition, currency, period, included and excluded records, source modules, record counts, as-of timestamp, and data-quality warnings.",
        "Missing, stale, unavailable, and unauthorized values remain distinct. Missing values are never silently converted to zero.",
    ]),
    ("11. Private Model and Response-Depth Architecture", [
        "The private model understands intent, selects approved tools and schemas, combines authorized evidence, explains deterministic calculations, compares periods, identifies drivers and anomalies, and produces structured drafts.",
        "The model does not override source authority. Permissions come from Pulse authorization; financial values come from deterministic calculations; project dates come from FlowHive; record state comes from the owning module.",
        "Analytical responses use a deep profile: direct conclusion, scope and filters, detailed evidence, calculations, conflicts, assumptions, limitations, risks, recommended actions, navigation, and freshness.",
    ]),
    ("12. Controlled External LLM Integration", [
        "External providers are optional generic reasoning resources. They are not the primary location for Pulse internal documents or restricted system context.",
        "The DLP gateway removes document text, identities, record IDs, secrets, URLs, IP addresses, infrastructure identifiers, pricing, rates, revenue, cost, margin, contract terms, and unnecessary customer information.",
        "Module 064 owns provider credentials, model allowlists, health, feature routes, rate limits, timeouts, retries, circuit breakers, refusal handling, and sanitized provider telemetry.",
        "A safety refusal terminates routing. Pulse must not try another provider or local template to bypass a refusal.",
        "External responses are untrusted suggestions until the private verifier confirms them against authoritative documents, tools, calculations, and policies.",
    ]),
    ("13. Security and Privacy Architecture", [
        "Authenticate every protected request and fail closed when identity, permission, project, document, or tool authorization evidence is unavailable.",
        "Resolve actual and effective View-As identities. View-As remains read-only and never transfers mutation authority.",
        "Encrypt data in transit and at rest; use managed identity or approved service credentials; never embed secrets in prompts, source code, logs, or browser telemetry.",
        "Protect against prompt injection by separating instructions from retrieved data, labeling source content, applying tool allowlists, validating outputs, and never executing commands found in documents.",
        "Keep raw documents, extracted text, embeddings, prompts, model outputs, and audit metadata under separate least-privilege and retention controls.",
    ]),
    ("14. Deployment and Network Architecture", [
        "The deployment architecture separates edge and identity, Pulse application services, private data services, private AI services, semantic tools, audit services, and controlled outbound egress.",
        "Private endpoints, network security groups, service identities, DNS controls, certificate management, and outbound allowlists enforce the trust boundaries.",
        "The private model and vector index do not require public inbound access. Only Pulse orchestration and approved administrative services may reach them.",
        "External provider traffic originates from the controlled egress gateway and Module 064, not from browsers, documents, or arbitrary application components.",
    ]),
    ("15. Availability, Performance, and Resilience", [
        "Use health checks, timeouts, retry limits, circuit breakers, bulkheads, and bounded queues for private and external model services.",
        "Keep AI failures independent from transactional Pulse workflows so time entry, project management, billing, and other core functions continue.",
        "Preserve deterministic local templates for limited degraded operation without labeling them as document-grounded AI.",
        "Maintain tested rollback targets for model version, prompt version, retrieval configuration, feature route, and provider route.",
    ]),
    ("16. Observability, Audit, and Evaluation", [
        "Capture correlation ID, feature, actual/effective identity, scope summary, tool plan, source IDs and versions, model/prompt version, provider path, latency, token use, confidence, verification result, and user acceptance or correction.",
        "Do not log raw document chunks, secrets, unrestricted prompts, or sensitive model responses.",
        "Track grounded-answer rate, citation coverage, unauthorized retrieval rate, hallucination rate, tool-selection accuracy, timesheet edit distance, FlowHive correction rate, financial calculation accuracy, latency, cost, and DLP redaction completeness.",
        "Unauthorized retrieval must remain zero in preproduction validation and production monitoring.",
    ]),
    ("17. Model Lifecycle and Training", [
        "Use RAG and live tools first so changing facts remain current, correctable, and revocable.",
        "Capture accepted and corrected examples only as training candidates. Sanitize, review, version, checksum, and approve every dataset.",
        "Fine-tune an approved open-weight base model with LoRA or QLoRA in a private training environment only when evaluations justify the investment.",
        "Register the base model, license, dataset version, parameters, code revision, artifact checksum, evaluation results, approvals, deployment route, and rollback target.",
        "Production promotion remains human approved and requires frozen functional, privacy, security, and regression suites.",
    ]),
    ("18. Self-Sustaining Operating Model", [
        "Celar AI may automatically detect approved document changes, refresh indexes, revoke inaccessible content, monitor source freshness, capture feedback, run scheduled evaluations, detect quality drift, and prepare sanitized training candidates.",
        "Humans approve new data sources, classifications, datasets, training jobs, base models, model versions, external escalation policies, feature assignments, and production promotion.",
        "Celar AI never autonomously modifies its own policies, tools, prompts, permissions, training data, model weights, deployment, or production route.",
    ]),
    ("19. Implementation Roadmap", [
        "Phase 1 — Foundation: Module 011 governance, Module 064 provider boundary, answer contract, privacy policy, and architecture approval.",
        "Phase 2 — Private document pipeline: scanning, extraction, OCR, classification, versioning, citations, private embeddings, and permission-filtered retrieval.",
        "Phase 3 — Timesheet and Help/Search: document-grounded suggestions, detailed product Help, and governed read-tool planning.",
        "Phase 4 — FlowHive and financial insight: private document-to-plan generation, deterministic scheduling, semantic financial tools, and cited analysis.",
        "Phase 5 — Private model lifecycle: evaluation, LoRA/QLoRA training, model registry, canary, rollback, and continuous quality operations.",
    ]),
    ("20. Risks and Mitigations", [
        "Information leakage — mitigate with authorization-first retrieval, DLP, private models, outbound allowlists, and zero raw-document external routing.",
        "Hallucination — mitigate with source-grounded generation, deterministic tools, citation requirements, confidence thresholds, and private verification.",
        "Prompt injection — mitigate with content labeling, instruction isolation, tool allowlists, output validation, and no command execution from documents.",
        "Stale or conflicting data — mitigate with version authority, freshness indicators, conflict surfacing, and source-owner remediation.",
        "Over-broad access — mitigate with actual/effective identity, module/action permission, project/customer/team/record scope, and field-level filtering.",
        "Operational dependence — mitigate with degraded templates, independent transactional workflows, capacity planning, observability, and rollback.",
    ]),
    ("21. Acceptance Criteria", [
        "Unauthorized document and record retrieval tests return zero content.",
        "Raw SOW, GSD, contract, customer, architecture, financial, and employee content is not sent to public external providers.",
        "Timesheet suggestions identify evidence and cannot change hours, save, submit, or approve.",
        "Help and Search answers include sources, scope, filters, as-of time, and clear uncertainty.",
        "FlowHive drafts include citations, assumptions, risks, dependencies, and unresolved questions and cannot baseline or commit dates.",
        "Financial answers use governed formulas and preserve unknown values.",
        "Every external capsule passes DLP and is privately verified before display.",
        "Model promotion requires approved datasets, passing evaluations, human approval, and tested rollback.",
    ]),
    ("22. Celar AI Identity, Origin, and Operating Narrative", [
        "Core identity: Celar AI is the unified operational intelligence system for the US Signal Solution Provider division. Conceived and engineered under the direction of Dr. Ahmed Adeyemi, Manager of Professional Services, it creates one governed intersection for Sales, Professional Services, Project Management, Engineering, Finance, Operations, Security, and leadership.",
        "Meaning behind the name: Celar AI draws from celeritas, the Latin concept of swiftness or speed, and from c, the conventional symbol for the speed of light in E=mc². The name connects US Signal's fiber-optic network foundation to a Professional Services promise of faster delivery.",
        "Speed-of-delivery mission: Celar AI reduces the elapsed time required to scope, prepare and review SOWs, hand work to delivery teams, plan projects, record time, resolve operational questions, understand financial health, support invoices, and close work.",
        "Changepoint catalyst: Changepoint provided a functional legacy PSA and system of record, but siloed data, rigid navigation, repetitive administration, and manual transfers created an administration tax and a speed bottleneck for consulting teams.",
        "Celar AI does not remove governance. It replaces fragmented administration with permission-aware retrieval, deterministic tools, private reasoning, complete audit evidence, and human-reviewed actions.",
        "Canonical response: Celar AI is the unified operational intelligence system for the US Signal Solution Provider division. It was conceived and engineered under the direction of Dr. Ahmed Adeyemi, Manager of Professional Services, to create a central intersection where consulting teams convene, collaborate, and exchange project, delivery, operational, and financial information. Its name draws from celeritas and the speed-of-light symbol c, reflecting US Signal's fiber heritage and its mission to turn speed of light into speed of delivery. The system addresses the operational drag associated with legacy Changepoint workflows by unifying authorized documents, data, workflows, and AI-assisted reasoning so teams can scope, execute, troubleshoot, report, and invoice work more quickly.",
        "Transition boundary: Version 2.0 establishes Celar AI as the target documentation brand. Existing Pulse AI runtime labels, routes, APIs, source directories, database objects, permissions, feature codes, environment variables, Module 064 settings, and deployed resources remain unchanged until a separate application-rebrand package is approved.",
        "Brand governance: Celar AI is strategically aligned with US Signal's fiber and solution-provider mission, but Celar is not globally unique. Public marketing, trademark filing, domain acquisition, or customer-facing launch requires formal US Signal Legal and Marketing clearance.",
    ]),
]

FEATURE_ROUTES = [
    ["Feature", "Primary path", "External policy", "Human control"],
    ["Timesheet — document grounded", "Private model + RAG", "No raw-document external route", "Engineer reviews and applies"],
    ["Timesheet — basic note", "Private model", "Sanitized external fallback allowed by policy", "Engineer reviews and applies"],
    ["Help and Search", "Private model + governed tools", "Sanitized generic reasoning only", "User receives sources and uncertainty"],
    ["FlowHive planning", "Private model + schedule engine", "Generic planning checklist only", "PM and Engineering modify before baseline"],
    ["Reports and financials", "Deterministic tools + private explanation", "Disabled by default", "No financial mutation"],
]

ADRS = [
    ["ADR", "Decision", "Rationale"],
    ["ADR-001", "Private-first model path", "Restricted Pulse data remains inside the approved trust boundary."],
    ["ADR-002", "Authorization before retrieval", "The application remains the access-control authority."],
    ["ADR-003", "Module 064 controls external providers", "One provider, secret, health, and routing boundary prevents bypass."],
    ["ADR-004", "Deterministic calculations", "Models explain financial and schedule results but do not invent them."],
    ["ADR-005", "Human-approved learning", "Feedback may become training data only after review and version approval."],
]

GLOSSARY = [
    ["Term", "Definition"],
    ["Pulse", "US Signal business platform and application boundary."],
    ["Celar AI", "Module 011 private intelligence, knowledge, evaluation, and model-lifecycle capability."],
    ["Module 064", "Governed external-provider configuration, health, routing, and fallback gateway."],
    ["RAG", "Retrieval-augmented generation using authorized source evidence."],
    ["DLP", "Data-loss-prevention controls that detect and remove sensitive content."],
    ["Reasoning capsule", "Minimal sanitized problem statement eligible for optional external reasoning."],
    ["Effective user", "The identity whose permissions and data scope apply, including read-only View-As."],
    ["Private verifier", "Service that re-grounds model output against documents, tools, calculations, and policy."],
]


def esc(value: str) -> str:
    return html.escape(str(value), quote=True)


def logo_data_uri() -> str:
    return "data:image/jpeg;base64," + base64.b64encode(LOGO.read_bytes()).decode("ascii")


def svg_header(width: int, height: int, title: str) -> list[str]:
    return [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<defs>',
        '<filter id="shadow" x="-20%" y="-20%" width="140%" height="140%"><feDropShadow dx="0" dy="4" stdDeviation="5" flood-color="#0B2540" flood-opacity="0.18"/></filter>',
        '<marker id="arrow" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="#1769AA"/></marker>',
        '<marker id="arrow-green" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="#67B446"/></marker>',
        '<style>text{font-family:Arial,Helvetica,sans-serif}.title{font-size:34px;font-weight:700;fill:#072D59}.sub{font-size:16px;fill:#465B6D}.zone{font-size:16px;font-weight:700;fill:#072D59;letter-spacing:.4px}.box-title{font-size:16px;font-weight:700;fill:#072D59}.box-copy{font-size:12px;fill:#34495E}.small{font-size:10px;fill:#51697C}.label{font-size:11px;font-weight:700;fill:#1769AA}</style>',
        '</defs>', f'<rect width="{width}" height="{height}" fill="#FFFFFF"/>',
        f'<image href="{logo_data_uri()}" x="40" y="24" width="150" height="95" preserveAspectRatio="xMidYMid meet"/>',
        f'<text x="220" y="62" class="title">{esc(title)}</text>',
        f'<text x="220" y="91" class="sub">US Signal • Pulse • Module 011 • Celeritas to speed of delivery • Version {VERSION}</text>',
        f'<line x1="40" y1="125" x2="{width-40}" y2="125" stroke="{BLUE}" stroke-width="4"/>',
    ]


def svg_box(x, y, w, h, title, lines: Sequence[str], fill=WHITE, stroke=MID_BLUE, radius=12, tag=None):
    out = [f'<g filter="url(#shadow)"><rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{radius}" fill="{fill}" stroke="{stroke}" stroke-width="2"/></g>']
    if tag:
        out.append(f'<rect x="{x+12}" y="{y+12}" width="{max(70, len(tag)*7)}" height="22" rx="11" fill="{stroke}"/><text x="{x+22}" y="{y+28}" font-size="11" font-weight="700" fill="#FFFFFF">{esc(tag)}</text>')
        title_y = y + 54
    else:
        title_y = y + 30
    out.append(f'<text x="{x+16}" y="{title_y}" class="box-title">{esc(title)}</text>')
    for i, line in enumerate(lines):
        out.append(f'<text x="{x+16}" y="{title_y+23+i*17}" class="box-copy">{esc(line)}</text>')
    return out


def arrow(x1, y1, x2, y2, label=None, dashed=False, green=False):
    color = GREEN if green else MID_BLUE
    marker = 'arrow-green' if green else 'arrow'
    dash = ' stroke-dasharray="8 6"' if dashed else ''
    out = [f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{color}" stroke-width="3" marker-end="url(#{marker})"{dash}/>' ]
    if label:
        out.append(f'<rect x="{(x1+x2)/2-75}" y="{(y1+y2)/2-13}" width="150" height="22" rx="6" fill="#FFFFFF" opacity="0.94"/><text x="{(x1+x2)/2}" y="{(y1+y2)/2+3}" text-anchor="middle" class="label">{esc(label)}</text>')
    return out


def build_logical_svg() -> str:
    w, h = 1600, 1050
    s = svg_header(w, h, "Celar AI Private-First Logical Architecture — Speed of Delivery")
    for x,y,zw,zh,label,fill in [(35,150,1530,165,"PULSE APPLICATION TRUST ZONE",LIGHT_BLUE),(35,330,1530,390,"PRIVATE DATA AND AI TRUST ZONE","#F7FAFC"),(35,735,930,250,"PRIVATE VERIFICATION AND RESPONSE ZONE","#F4F9F1"),(985,735,580,250,"CONTROLLED EGRESS / EXTERNAL PROVIDER ZONE","#FFF8EE")]:
        s += [f'<rect x="{x}" y="{y}" width="{zw}" height="{zh}" rx="16" fill="{fill}" stroke="#9BB2C6" stroke-width="1.5"/>', f'<text x="{x+18}" y="{y+28}" class="zone">{label}</text>']
    s += svg_box(75,200,250,90,"Pulse User Experience",["Timesheet • Ask Celar AI • Search","FlowHive • Reports • Financials"],tag="ENTRY")
    s += svg_box(390,200,300,90,"Identity and Authorization",["Session • actual/effective user","RBAC • module • project • record scope"])
    s += svg_box(760,200,300,90,"Celar AI Orchestrator",["Intent • policy • tool plan","minimum-necessary context assembly"])
    s += svg_box(1130,200,360,90,"Audit and Correlation",["Request ID • scope summary • evidence","sanitized telemetry only"])
    s += arrow(325,245,390,245)+arrow(690,245,760,245)+arrow(1060,245,1130,245,dashed=True)
    s += svg_box(75,385,300,120,"Private Document Retrieval",["SOW • GSD • design • contract","versioned extraction and citations","permission-filtered hybrid/vector search"],tag="RAG")
    s += svg_box(425,385,300,120,"Governed Live-Data Tools",["projects • tasks • time • approvals","capacity • reports • finance","allowlisted read-only semantic contracts"],tag="TOOLS")
    s += svg_box(785,385,300,120,"Private Celar AI Model",["open-weight private endpoint","structured deep-response schemas","no public inbound access"],fill="#EDF8FE",stroke=BLUE,tag="PRIMARY")
    s += svg_box(1145,385,345,120,"Confidence Assessment",["coverage • freshness • conflicts","calculation validity • unsupported claims","external-help policy eligibility"],fill="#F1F8ED",stroke=GREEN,tag="GATE")
    s += arrow(910,290,225,385,"authorized document plan")+arrow(910,290,575,385,"governed tool plan")+arrow(375,445,785,445,"cited context")+arrow(725,465,785,465,"deterministic results")+arrow(1085,445,1145,445)
    s += svg_box(75,565,300,110,"Private Storage and Index",["object storage • PostgreSQL","extracted text • embeddings","retention and access metadata"])
    s += svg_box(425,565,300,110,"Prompt and Schema Registry",["feature prompts • response schemas","versioning • approval • rollback"])
    s += svg_box(785,565,300,110,"Model Lifecycle Services",["evaluation • feedback • registry","private training • canary • rollback"])
    s += svg_box(1145,565,345,110,"Policy and DLP Controls",["classification • redaction","egress allowlists • refusal policy"])
    s += arrow(225,505,225,565,dashed=True)+arrow(575,505,575,565,dashed=True)+arrow(935,505,935,565,dashed=True)+arrow(1318,505,1318,565,dashed=True)
    s += svg_box(75,790,300,125,"Sufficient Evidence",["complete private answer path","no external provider required","preserve sources and uncertainty"],fill="#F1F8ED",stroke=GREEN,tag="PASS")
    s += svg_box(425,790,300,125,"Private Evidence Verification",["re-ground claims and calculations","remove unsupported statements","attach citations and navigation"],tag="VERIFY")
    s += svg_box(785,790,140,125,"Detailed Result",["cited answer","or reviewable draft"],fill="#EDF8FE",stroke=BLUE)
    s += arrow(1318,505,225,790,"confidence sufficient",green=True)+arrow(375,852,425,852,green=True)+arrow(725,852,785,852,green=True)
    s += svg_box(1025,790,220,125,"Sanitization / DLP",["minimal generic problem","no raw documents, IDs,","secrets, pricing, or PII"],fill="#FFFDF8",stroke="#D9992E",tag="OPTIONAL")
    s += svg_box(1285,790,120,125,"Module 064",["provider","health","routing"],stroke="#D9992E")
    s += svg_box(1440,790,90,125,"Claude /",["OpenAI","generic","reasoning"],stroke="#D9992E")
    s += arrow(1318,505,1135,790,"generic help eligible",dashed=True)+arrow(1245,852,1285,852,dashed=True)+arrow(1405,852,1440,852,dashed=True)+arrow(1485,915,700,915,"untrusted suggestion returned for private verification",dashed=True)
    s += [f'<text x="40" y="1025" class="small">{CLASSIFICATION} • Architecture baseline v{VERSION} • Raw internal documents remain inside the private Pulse boundary by default.</text>','</svg>']
    return "\n".join(s)


def build_deployment_svg() -> str:
    w, h = 1800, 1100
    s = svg_header(w, h, "Celar AI Deployment and Network Architecture")
    for x,y,zw,zh,label,fill in [(35,150,300,875,"EDGE AND IDENTITY",LIGHT_BLUE),(355,150,390,875,"PULSE APPLICATION ZONE","#F7FAFC"),(765,150,410,875,"PRIVATE DATA SERVICES","#F7FAFC"),(1195,150,330,875,"PRIVATE AI SERVICES","#F4F9F1"),(1545,150,220,875,"CONTROLLED EGRESS","#FFF8EE")]:
        s += [f'<rect x="{x}" y="{y}" width="{zw}" height="{zh}" rx="16" fill="{fill}" stroke="#9BB2C6" stroke-width="1.5"/>',f'<text x="{x+18}" y="{y+28}" class="zone">{label}</text>']
    s += svg_box(70,210,230,100,"Users and Devices",["managed browsers","Pulse module experiences"])+svg_box(70,350,230,110,"Edge Security",["DNS • TLS • WAF","reverse proxy • rate limits"])+svg_box(70,505,230,115,"Identity Provider",["SSO • MFA • conditional access","session identity claims"])+svg_box(70,665,230,120,"Administrative Access",["approved operator roles","private management path","no direct model public access"])
    s += svg_box(390,210,320,115,"Pulse Web and API",["authenticated application shell","module endpoints • request correlation","transactional workflows remain independent"],tag="PULSE")+svg_box(390,370,320,120,"Authorization and Policy",["actual/effective user • RBAC","module/action/project/customer scope","field and classification controls"])+svg_box(390,535,320,120,"Celar AI Orchestrator",["intent • retrieval • tool plan","prompt/schema selection","confidence and response composition"],fill="#EDF8FE",stroke=BLUE)+svg_box(390,700,320,120,"Governed Tool Gateway",["allowlisted read-only APIs","semantic metrics • row limits","source health and freshness"])+svg_box(390,865,320,105,"Audit and Observability",["correlation • evidence • quality","sanitized logs • alerts • SLOs"])
    s += svg_box(800,210,340,115,"Private Object Storage",["original documents • extracted artifacts","versioning • malware scan • retention"])+svg_box(800,370,340,115,"Pulse PostgreSQL",["authoritative metadata and records","approvals • registry • audit evidence"])+svg_box(800,530,340,120,"Extraction and OCR Workers",["PDF • DOCX • XLSX • text","OCR only when required","classification • chunking • citations"])+svg_box(800,695,340,120,"Permission-Scoped Search",["keyword + semantic/vector index","security filters before ranking","revocation and freshness controls"])+svg_box(800,860,340,110,"Key and Secret Services",["managed identities • key management","write-only provider secrets"])
    s += svg_box(1230,210,260,125,"Private Embedding Service",["private endpoint","approved embedding model","no public data path"],stroke=GREEN)+svg_box(1230,380,260,135,"Private Celar AI Model",["open-weight inference endpoint","restricted-context reasoning","structured output • citations"],fill="#F1F8ED",stroke=GREEN,tag="PRIMARY")+svg_box(1230,560,260,120,"Evaluation and Feedback",["frozen test suites","accept/edit/reject evidence","drift and regression monitoring"],stroke=GREEN)+svg_box(1230,725,260,120,"Model Registry and Training",["dataset approval • LoRA/QLoRA","artifact checksums • canary","human-approved promotion"],stroke=GREEN)+svg_box(1230,890,260,80,"Private Verifier",["re-ground claims • enforce policy"],stroke=GREEN)
    s += svg_box(1575,210,160,130,"DLP / Redaction",["classify","minimize","sanitize","block"],stroke="#D9992E")+svg_box(1575,390,160,130,"Module 064",["credentials","model allowlists","health • routes","circuit breakers"],stroke="#D9992E")+svg_box(1575,570,160,130,"External LLM",["Claude / OpenAI","generic reasoning","sanitized capsule","optional only"],stroke="#D9992E")+svg_box(1575,750,160,110,"External Telemetry",["sanitized status","latency • usage","no restricted content"],stroke="#D9992E")
    s += arrow(185,310,185,350)+arrow(300,405,390,267,"TLS")+arrow(300,565,390,425,"identity claims")+arrow(550,325,550,370)+arrow(550,490,550,535)+arrow(710,595,800,755,"authorized retrieval")+arrow(710,760,800,755,"tool data")+arrow(970,650,970,695)+arrow(1140,755,1230,445,"cited context")+arrow(1140,275,1230,275,"content for embedding",dashed=True)+arrow(1360,515,1360,560)+arrow(1360,680,1360,725,dashed=True)+arrow(1360,845,1360,890,dashed=True)+arrow(1490,445,1575,275,"sanitized only",dashed=True)+arrow(1655,340,1655,390,dashed=True)+arrow(1655,520,1655,570,dashed=True)+arrow(1575,635,1490,930,"untrusted result",dashed=True)+arrow(1230,930,710,595,"verified response",green=True)+arrow(550,820,550,865,dashed=True)+arrow(970,970,550,920,"audit events",dashed=True)+arrow(1360,970,710,920,"model telemetry",dashed=True)+arrow(1655,860,710,920,"provider telemetry",dashed=True)
    s += [f'<text x="40" y="1070" class="small">{CLASSIFICATION} • Private endpoints and outbound allowlists enforce the trust zones. External providers never receive unrestricted Pulse context.</text>','</svg>']
    return "\n".join(s)


def write_diagrams():
    logical = OUT / FILES["logical_svg"]
    deployment = OUT / FILES["deployment_svg"]
    logical.write_text(build_logical_svg(), encoding="utf-8")
    deployment.write_text(build_deployment_svg(), encoding="utf-8")
    cairosvg.svg2png(bytestring=logical.read_bytes(), write_to=str(OUT / FILES["logical_png"]), output_width=2400)
    cairosvg.svg2png(bytestring=deployment.read_bytes(), write_to=str(OUT / FILES["deployment_png"]), output_width=2700)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_text_color(cell, color_hex: str):
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))


def add_docx_table(doc: Document, rows: Sequence[Sequence[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0])); table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r,c); cell.text = str(value); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Arial"; run.font.size = Pt(8.5); run.bold = r == 0
            if r == 0: set_cell_shading(cell,NAVY); set_cell_text_color(cell,WHITE)
            elif r % 2 == 0: set_cell_shading(cell,"F2F5F7")
    doc.add_paragraph(); return table


def configure_docx(doc: Document):
    normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(9.5); normal.font.color.rgb=RGBColor.from_string(BLACK.replace("#","")); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.08
    for name,size,color in [("Title",28,NAVY),("Heading 1",18,NAVY),("Heading 2",13,MID_BLUE),("Heading 3",10.5,DARK_GRAY)]:
        st=doc.styles[name]; st.font.name="Arial"; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color.replace("#","")); st.paragraph_format.space_before=Pt(8); st.paragraph_format.space_after=Pt(5)
    doc.styles["List Bullet"].font.name="Arial"; doc.styles["List Bullet"].font.size=Pt(9.2)
    for section in doc.sections:
        section.top_margin=Inches(.55); section.bottom_margin=Inches(.55); section.left_margin=Inches(.7); section.right_margin=Inches(.7)
        hp=section.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=hp.add_run(f"US Signal  |  Celar AI Architecture  |  v{VERSION}"); r.font.name="Arial"; r.font.size=Pt(7.5); r.font.color.rgb=RGBColor.from_string(MID_BLUE.replace("#",""))
        fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=fp.add_run(CLASSIFICATION+"   |   "); r.font.name="Arial"; r.font.size=Pt(7); r.font.color.rgb=RGBColor.from_string(DARK_GRAY.replace("#","")); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)


def build_docx():
    doc=Document(); configure_docx(doc)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(LOGO),width=Inches(1.6))
    p=doc.add_paragraph("US SIGNAL CREATED ARCHITECTURE DOCUMENT"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.runs[0]; r.font.name="Arial"; r.font.size=Pt(10); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE.replace('#',''))
    p=doc.add_paragraph(DOC_TITLE,style="Title"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph("Celeritas: speed of light. Celar AI: speed of delivery. Private-first intelligence for unified solution-provider operations"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.size=Pt(13); p.runs[0].font.color.rgb=RGBColor.from_string(DARK_GRAY.replace('#',''))
    doc.add_paragraph(); add_docx_table(doc,[["Document owner","Product","Version","Classification"],[OWNER,PRODUCT,VERSION,CLASSIFICATION],["Architecture status","Module","Prepared for","Canonical name"],["Review baseline","011 — Celar AI","US Signal architecture review","Pulse"]])
    p=doc.add_paragraph("This architecture package is created and owned by US Signal. It uses the exact US Signal logo supplied for the package and refers to the application as Pulse."); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; doc.add_page_break()
    doc.add_heading("Document Control",level=1); add_docx_table(doc,[["Field","Value"],["Title",DOC_TITLE],["Owner",OWNER],["Product",PRODUCT],["Module","011 — Celar AI"],["Version",VERSION],["Status","Review baseline"],["Classification",CLASSIFICATION],["Purpose","Explain the private-first Celar AI architecture to systems engineering, security, application, data, AI, project, financial, and leadership stakeholders."]])
    doc.add_heading("Revision History",level=2); add_docx_table(doc,[["Version","Date","Description"],["1.0","2026-07-29","Initial architecture package."],["1.1","2026-07-29","Corrected to the supplied US Signal logo, US Signal ownership, Pulse naming, and US Signal visual direction."]]); doc.add_heading("Naming Convention",level=2); doc.add_paragraph("Pulse is the business platform. Celar AI is the private intelligence capability in Module 011. Module 064 is the governed external-provider configuration, health, routing, circuit-breaker, and fallback gateway."); doc.add_page_break()
    doc.add_heading("Contents",level=1)
    for title,_ in SECTIONS: doc.add_paragraph(title)
    for appendix in ["Appendix A. Feature Routing Matrix","Appendix B. Architecture Decision Records","Appendix C. Glossary"]: doc.add_paragraph(appendix)
    doc.add_page_break()
    for title,paragraphs in SECTIONS:
        doc.add_heading(title,level=1)
        if title.startswith("4."): doc.add_paragraph("Figure 1. Celar AI private-first logical architecture and confidence-driven external escalation path."); doc.add_picture(str(OUT/FILES["logical_png"]),width=Inches(7.0))
        if title.startswith("14."): doc.add_paragraph("Figure 2. Celar AI deployment zones, private data paths, governed tools, controlled egress, and audit integration."); doc.add_picture(str(OUT/FILES["deployment_png"]),width=Inches(7.0))
        for text in paragraphs: doc.add_paragraph(text,style="List Bullet" if text.startswith(("Private first:","Authorization before","Facts from","Progressive","External models","Human control:")) else "Normal")
        if title.startswith("5."): add_docx_table(doc,[["Layer","Key components","Authority"],["Experience","Timesheet, Ask Celar AI, Search, FlowHive, Reports","Owning Pulse modules"],["Policy","Identity, RBAC, scope, classification","Pulse backend"],["Knowledge","Storage, extraction, embeddings, search","Private data services"],["Tools","Read-only APIs and semantic metrics","Owning modules"],["Reasoning","Private model, confidence, verifier","Module 011"],["Egress","DLP, Module 064, external providers","Security policy + Module 064"]])
        if title.startswith("6."): add_docx_table(doc,[["Data class","Private model","External provider"],["Public operating guidance","Allowed","Allowed when routed"],["Internal module documentation","Allowed","Sanitized generic question only"],["SOW, GSD, architecture and customer documents","Allowed","Raw content prohibited"],["Rates, contracts and financial data","Allowed","Disabled by default"],["Credentials and secrets","Prohibited in prompts","Prohibited"]])
        if title.startswith("10."): add_docx_table(doc,[["Semantic element","Example"],["Metrics","planned cost, actual cost, forecast, variance, margin"],["Dimensions","project, customer, project manager, period"],["Filters","effective user, authorized workspace, date range, status"],["Rules","read-only, deterministic values, unknowns preserved, row limit"]])
        if title.startswith("13."): add_docx_table(doc,[["Threat","Required mitigation"],["Prompt injection","Instruction isolation, content labels, tool allowlists, output validation"],["Data leakage","Authorization-first retrieval, DLP, private model, outbound allowlists"],["Hallucination","Citations, deterministic tools, confidence gates, private verification"],["Privilege escalation","Actual/effective identity, server-side scope, View-As read-only"],["Model drift","Frozen evaluations, canary, monitoring, rollback"]])
        if title.startswith("19."): add_docx_table(doc,[["Phase","Deliverable","Gate"],["1","Foundation and architecture","Security and architecture approval"],["2","Private extraction and RAG","Privacy and retrieval tests"],["3","Timesheet and Help/Search","User acceptance and grounding quality"],["4","FlowHive and financial insight","Deterministic tool integration"],["5","Training and model lifecycle","Evaluation and production approval"]])
        doc.add_page_break()
    doc.add_heading("Appendix A. Feature Routing Matrix",level=1); add_docx_table(doc,FEATURE_ROUTES); doc.add_page_break(); doc.add_heading("Appendix B. Architecture Decision Records",level=1); add_docx_table(doc,ADRS); doc.add_page_break(); doc.add_heading("Appendix C. Glossary",level=1); add_docx_table(doc,GLOSSARY)
    doc.save(OUT/FILES["docx"])


def pdf_styles():
    base=getSampleStyleSheet()
    return {"title":ParagraphStyle("TitleUS",parent=base["Title"],fontName="Helvetica-Bold",fontSize=27,leading=31,textColor=colors.HexColor(NAVY),alignment=TA_CENTER,spaceAfter=14),"subtitle":ParagraphStyle("SubtitleUS",parent=base["BodyText"],fontName="Helvetica",fontSize=12.5,leading=17,textColor=colors.HexColor(DARK_GRAY),alignment=TA_CENTER,spaceAfter=14),"h1":ParagraphStyle("H1US",parent=base["Heading1"],fontName="Helvetica-Bold",fontSize=18,leading=22,textColor=colors.HexColor(NAVY),spaceAfter=10),"body":ParagraphStyle("BodyUS",parent=base["BodyText"],fontName="Helvetica",fontSize=9.2,leading=13,textColor=colors.HexColor(BLACK),spaceAfter=7),"bullet":ParagraphStyle("BulletUS",parent=base["BodyText"],fontName="Helvetica",fontSize=9,leading=12.5,leftIndent=14,firstLineIndent=-8,bulletIndent=5,textColor=colors.HexColor(BLACK),spaceAfter=4),"small":ParagraphStyle("SmallUS",parent=base["BodyText"],fontName="Helvetica",fontSize=7.5,leading=10,textColor=colors.HexColor(DARK_GRAY)),"caption":ParagraphStyle("CaptionUS",parent=base["BodyText"],fontName="Helvetica-Oblique",fontSize=8,leading=10,textColor=colors.HexColor(DARK_GRAY),alignment=TA_CENTER,spaceAfter=8)}


class NumberedDocTemplate(BaseDocTemplate):
    def __init__(self,filename,**kw):
        super().__init__(filename,**kw); self.addPageTemplates(PageTemplate(id="main",frames=Frame(self.leftMargin,self.bottomMargin,self.width,self.height,id="normal"),onPage=self._header_footer))
    def _header_footer(self,c,doc):
        c.saveState(); c.setStrokeColor(colors.HexColor(BLUE)); c.setLineWidth(1); c.line(doc.leftMargin,LETTER[1]-.47*inch,LETTER[0]-doc.rightMargin,LETTER[1]-.47*inch); c.setFont("Helvetica",7.5); c.setFillColor(colors.HexColor(MID_BLUE)); c.drawRightString(LETTER[0]-doc.rightMargin,LETTER[1]-.36*inch,f"US Signal | Celar AI Architecture | v{VERSION}"); c.setFillColor(colors.HexColor(DARK_GRAY)); c.setFont("Helvetica",7); c.drawString(doc.leftMargin,.34*inch,CLASSIFICATION); c.drawRightString(LETTER[0]-doc.rightMargin,.34*inch,f"Page {doc.page}"); c.restoreState()


def rl_table(rows:Sequence[Sequence[str]],widths=None):
    styles=pdf_styles(); head=ParagraphStyle("TableHeaderUS",parent=styles["small"],fontName="Helvetica-Bold",textColor=colors.white); data=[]
    for r,row in enumerate(rows): data.append([Paragraph(esc(v),head if r==0 else styles["small"]) for v in row])
    table=Table(data,colWidths=widths,repeatRows=1,hAlign="LEFT"); table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor(NAVY)),("GRID",(0,0),(-1,-1),.5,colors.HexColor(MED_GRAY)),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)]))
    for r in range(1,len(rows)):
        if r%2==0: table.setStyle(TableStyle([("BACKGROUND",(0,r),(-1,r),colors.HexColor(LIGHT_GRAY))]))
    return table


def build_pdf():
    st=pdf_styles(); doc=NumberedDocTemplate(str(OUT/FILES["pdf"]),pagesize=LETTER,rightMargin=.62*inch,leftMargin=.62*inch,topMargin=.65*inch,bottomMargin=.58*inch,title=DOC_TITLE,author=OWNER); story=[]
    story += [Spacer(1,.35*inch),RLImage(str(LOGO),width=1.55*inch,height=1.0*inch),Spacer(1,.1*inch),Paragraph("US SIGNAL CREATED ARCHITECTURE DOCUMENT",ParagraphStyle("eye",parent=st["small"],fontName="Helvetica-Bold",fontSize=10,textColor=colors.HexColor(BLUE),alignment=TA_CENTER)),Spacer(1,.15*inch),Paragraph(DOC_TITLE,st["title"]),Paragraph("Celeritas: speed of light. Celar AI: speed of delivery. Private-first intelligence for unified solution-provider operations",st["subtitle"]),Spacer(1,.2*inch),rl_table([["Document owner","Product","Version","Classification"],[OWNER,PRODUCT,VERSION,CLASSIFICATION],["Architecture status","Module","Prepared for","Canonical name"],["Review baseline","011 — Celar AI","US Signal architecture review","Pulse"]],[1.25*inch,1.05*inch,.7*inch,3.7*inch]),Spacer(1,.25*inch),Paragraph("This architecture package is created and owned by US Signal. Celar AI was conceived and engineered under the direction of Dr. Ahmed Adeyemi, Manager of Professional Services, and uses the supplied US Signal logo.",st["subtitle"]),PageBreak()]
    story += [Paragraph("Document Control",st["h1"]),rl_table([["Field","Value"],["Title",DOC_TITLE],["Owner",OWNER],["Product",PRODUCT],["Module","011 — Celar AI"],["Version",VERSION],["Status","Review baseline"],["Classification",CLASSIFICATION]],[1.4*inch,5.8*inch]),PageBreak(),Paragraph("Contents",st["h1"])]
    story += [Paragraph(esc(t),st["body"]) for t,_ in SECTIONS]+[Paragraph("Appendix A. Feature Routing Matrix",st["body"]),Paragraph("Appendix B. Architecture Decision Records",st["body"]),Paragraph("Appendix C. Glossary",st["body"]),PageBreak()]
    for title,paragraphs in SECTIONS:
        story.append(Paragraph(esc(title),st["h1"]))
        if title.startswith("4."): story += [RLImage(str(OUT/FILES["logical_png"]),width=7*inch,height=3.05*inch),Paragraph("Figure 1. Celar AI private-first logical architecture.",st["caption"])]
        if title.startswith("14."): story += [RLImage(str(OUT/FILES["deployment_png"]),width=7*inch,height=3.15*inch),Paragraph("Figure 2. Celar AI deployment and network architecture.",st["caption"])]
        for p in paragraphs: story.append(Paragraph(esc(p),st["bullet"] if p.startswith(("Private first:","Authorization before","Facts from","Progressive","External models","Human control:")) else st["body"],bulletText="•" if p.startswith(("Private first:","Authorization before","Facts from","Progressive","External models","Human control:")) else None))
        if title.startswith("5."): story.append(rl_table([["Layer","Key components","Authority"],["Experience","Timesheet, Ask Celar AI, Search, FlowHive, Reports","Owning modules"],["Policy","Identity, RBAC, scope, classification","Pulse backend"],["Knowledge","Storage, extraction, embeddings, search","Private data services"],["Tools","Read-only APIs and metrics","Owning modules"],["Reasoning","Private model, confidence, verifier","Module 011"],["Egress","DLP, Module 064, external providers","Security + Module 064"]],[1*inch,3.3*inch,2.9*inch]))
        if title.startswith("6."): story.append(rl_table([["Data class","Private model","External provider"],["Public guidance","Allowed","Allowed when routed"],["Internal documentation","Allowed","Sanitized question only"],["SOW/GSD/customer documents","Allowed","Raw content prohibited"],["Rates and financial data","Allowed","Disabled by default"],["Credentials and secrets","Prohibited in prompts","Prohibited"]],[2.4*inch,1.9*inch,2.9*inch]))
        if title.startswith("10."): story.append(rl_table([["Semantic element","Example"],["Metrics","planned cost, actual cost, forecast, variance, margin"],["Dimensions","project, customer, project manager, period"],["Filters","effective user, workspace, date range, status"],["Rules","read-only, deterministic, unknowns preserved, row limit"]],[1.6*inch,5.6*inch]))
        if title.startswith("13."): story.append(rl_table([["Threat","Required mitigation"],["Prompt injection","Instruction isolation, labels, tool allowlists, output validation"],["Data leakage","Authorization-first retrieval, DLP, private model, outbound allowlists"],["Hallucination","Citations, tools, confidence gates, private verification"],["Privilege escalation","Server-side actual/effective identity and scope"],["Model drift","Frozen evaluations, canary, monitoring, rollback"]],[1.6*inch,5.6*inch]))
        if title.startswith("19."): story.append(rl_table([["Phase","Deliverable","Gate"],["1","Foundation and architecture","Security and architecture approval"],["2","Private extraction and RAG","Privacy and retrieval tests"],["3","Timesheet and Help/Search","Grounding quality and UAT"],["4","FlowHive and financial insight","Deterministic tool integration"],["5","Training and model lifecycle","Evaluation and production approval"]],[.65*inch,3.3*inch,3.25*inch]))
        story.append(PageBreak())
    story += [Paragraph("Appendix A. Feature Routing Matrix",st["h1"]),rl_table(FEATURE_ROUTES,[1.45*inch,1.8*inch,2.2*inch,1.75*inch]),PageBreak(),Paragraph("Appendix B. Architecture Decision Records",st["h1"]),rl_table(ADRS,[.75*inch,2*inch,4.45*inch]),PageBreak(),Paragraph("Appendix C. Glossary",st["h1"]),rl_table(GLOSSARY,[1.45*inch,5.75*inch])]
    doc.build(story)


def build_diagram_pdf():
    page=landscape(A3); c=canvas.Canvas(str(OUT/FILES["diagrams_pdf"]),pagesize=page)
    for title,filename in [("Celar AI Private-First Logical Architecture — Speed of Delivery",FILES["logical_png"]),("Celar AI Deployment and Network Architecture",FILES["deployment_png"])]:
        c.setFillColor(colors.white); c.rect(0,0,page[0],page[1],fill=1,stroke=0); c.setFillColor(colors.HexColor(NAVY)); c.setFont("Helvetica-Bold",17); c.drawString(.45*inch,page[1]-.42*inch,title); c.setFillColor(colors.HexColor(DARK_GRAY)); c.setFont("Helvetica",8); c.drawRightString(page[0]-.45*inch,page[1]-.40*inch,f"US Signal • Pulse • v{VERSION}")
        img=Image.open(OUT/filename); iw,ih=img.size; scale=min((page[0]-.7*inch)/iw,(page[1]-.9*inch)/ih); w,h=iw*scale,ih*scale; c.drawImage(str(OUT/filename),(page[0]-w)/2,(page[1]-h)/2-.08*inch,width=w,height=h,preserveAspectRatio=True,mask='auto'); c.setFont("Helvetica",7); c.drawString(.45*inch,.25*inch,CLASSIFICATION); c.showPage()
    c.save()


def sha256(path:Path)->str:
    h=hashlib.sha256()
    with path.open('rb') as f:
        for chunk in iter(lambda:f.read(1024*1024),b''): h.update(chunk)
    return h.hexdigest()


def update_readme(hashes:dict[str,str]):
    readme=OUT/"README.md"
    if not readme.exists(): return
    text=readme.read_text(encoding='utf-8')
    for filename,digest in hashes.items(): text=re.sub(rf"(\| `{re.escape(filename)}` \|[^\n|]*\| `)[0-9a-f]{{64}}(` \|)",rf"\g<1>{digest}\2",text)
    text=re.sub(r"(\| `assets/US_Signal_Logo\.jpg` \|[^\n|]*\| `)[0-9a-f]{64}(` \|)",rf"\g<1>{sha256(LOGO)}\2",text)
    readme.write_text(text,encoding='utf-8')


def main():
    if not LOGO.exists(): raise SystemExit(f"Missing supplied US Signal logo: {LOGO}")
    write_diagrams(); build_docx(); build_pdf(); build_diagram_pdf()
    hashes={name:sha256(OUT/name) for name in FILES.values()}
    (OUT/"SHA256SUMS.txt").write_text("\n".join(f"{digest}  {name}" for name,digest in sorted(hashes.items()))+"\n",encoding='utf-8'); update_readme(hashes)
    print(f"CELAR_AI_ARCHITECTURE_PACKAGE_VERSION={VERSION}"); print(f"PULSE_AI_ARCHITECTURE_FILES={len(hashes)}")
    for name,digest in sorted(hashes.items()): print(f"{digest}  {name}")


if __name__ == "__main__": main()
